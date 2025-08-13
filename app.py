import streamlit as st
import re
import io
import fitz  # PyMuPDF
import docx
import nltk
import matplotlib.pyplot as plt
import seaborn as sns
from collections import Counter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

# Download NLTK data
nltk.download('punkt')
nltk.download('punkt_tab')

# ------------------------------
# Helper Functions
# ------------------------------

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in pdf_document:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    doc = docx.Document(docx_file)
    return "\n".join([para.text for para in doc.paragraphs])

def preprocess_text(text):
    """Clean and split into sentences"""
    text = re.sub(r'\s+', ' ', text).strip()
    sentences = nltk.sent_tokenize(text)
    return sentences

def extract_keywords(text):
    """Extract simple keywords (basic placeholder)"""
    words = re.findall(r'\b\w+\b', text.lower())
    stopwords = set(nltk.corpus.stopwords.words('english'))
    keywords = [w for w in words if w not in stopwords and len(w) > 2]
    return set(keywords)

def enhance_bullet_points(sentences, jd_keywords):
    """Match keywords and rank sentences by relevance"""
    results = []
    for sent in sentences:
        match_count = sum(1 for kw in jd_keywords if kw in sent.lower())
        if match_count > 0:
            results.append((sent.strip(), match_count))
    results.sort(key=lambda x: x[1], reverse=True)
    return results

def make_pdf(bullets):
    """Generate PDF with bullet points"""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Optimized Resume Bullets", styles['Heading1']))
    story.append(Spacer(1, 12))
    bullet_items = []
    for b in bullets:
        bullet_items.append(ListItem(Paragraph(b, styles['Normal'])))
    story.append(ListFlowable(bullet_items, bulletType='bullet', bulletColor=colors.HexColor("#000000")))
    doc.build(story)
    buf.seek(0)
    return buf.read()

def make_docx_bytes(bullets):
    """Generate DOCX with bullet points"""
    try:
        output = io.BytesIO()
        doc = docx.Document()
        doc.add_heading("Optimized Resume Bullets", level=1)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
        doc.save(output)
        output.seek(0)
        return output.read()
    except Exception:
        return None

def keyword_density_plot(resume_text, jd_keywords):
    """Generate keyword density heatmap"""
    words = re.findall(r'\b\w+\b', resume_text.lower())
    freq = Counter(words)
    density = {kw: freq.get(kw, 0) for kw in jd_keywords}
    density_sorted = dict(sorted(density.items(), key=lambda x: x[1], reverse=True))

    plt.figure(figsize=(8, 6))
    sns.heatmap(
        [[v] for v in density_sorted.values()],
        annot=True, fmt="d", cmap="YlGnBu",
        yticklabels=list(density_sorted.keys()), cbar=False
    )
    plt.title("Keyword Density Heatmap")
    st.pyplot(plt)

# ------------------------------
# Streamlit App
# ------------------------------

st.set_page_config(page_title="ATS Resume Optimizer", layout="wide")
st.title("📄 ATS Resume Optimizer")
st.markdown("Upload your Resume & Job Description to get **ATS-optimized bullet points** with keyword density analysis.")

# File Upload
col1, col2 = st.columns(2)
with col1:
    resume_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
with col2:
    jd_file = st.file_uploader("Upload Job Description (PDF/DOCX)", type=["pdf", "docx"])

if resume_file and jd_file:
    # Extract Resume Text
    if resume_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(resume_file)
    else:
        resume_text = extract_text_from_docx(resume_file)

    # Extract JD Text
    if jd_file.type == "application/pdf":
        jd_text = extract_text_from_pdf(jd_file)
    else:
        jd_text = extract_text_from_docx(jd_file)

    # Preprocess & Match
    resume_sentences = preprocess_text(resume_text)
    jd_keywords = extract_keywords(jd_text)
    ranked = enhance_bullet_points(resume_sentences, jd_keywords)

    if ranked:
        st.subheader("✅ Optimized Resume Bullet Points")
        for bullet, score in ranked:
            st.write(f"- {bullet}  *(match score: {score})*")

        # Heatmap
        st.subheader("📊 Keyword Density Heatmap")
        keyword_density_plot(resume_text, jd_keywords)

        # Downloads
        ordered_plain = [r[0] for r in ranked]

        pdf_bytes = make_pdf(ordered_plain)
        st.download_button(
            "📥 Download as PDF",
            data=pdf_bytes,
            file_name="optimized_resume_bullets.pdf",
            mime="application/pdf"
        )

        docx_bytes = make_docx_bytes(ordered_plain)
        if docx_bytes:
            st.download_button(
                "📥 Download as DOCX",
                data=docx_bytes,
                file_name="optimized_resume_bullets.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:
        st.warning("No matching keywords found between Resume and Job Description.")

st.markdown("---")
st.markdown("⚙ **Modular Design:** This code is structured so it can integrate directly with **Task 3** for automated company role analysis.")
